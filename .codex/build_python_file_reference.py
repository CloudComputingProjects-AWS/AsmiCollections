from __future__ import annotations

import ast
from collections import defaultdict
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION_START
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path.cwd()
DATE_LABEL = "14 May 2026"

BLUE = RGBColor(46, 116, 181)
NAVY = RGBColor(11, 37, 69)
DARK = RGBColor(34, 34, 34)
MUTED = RGBColor(90, 90, 90)
LIGHT_FILL = "F2F4F7"
CALLOUT_FILL = "F8FAFD"


OVERRIDES = {
    "backend/lambda_handler.py": "Primary AWS Lambda entrypoint for the FastAPI backend. Lazily loads the ASGI app for HTTP traffic and also routes scheduled task payloads for reservation release, FX sync, and deletion processing.",
    "backend/app/main.py": "Builds the FastAPI application used in local development and Lambda. Registers lifespan startup or shutdown behavior, middleware, health check, and the routed API surface.",
    "backend/app/api/v1/router.py": "Central router registry for the versioned API. Imports and mounts each domain router so app.main can expose a single /api/v1 surface.",
    "backend/app/core/config.py": "Holds the environment-driven settings model. It is the canonical source for application, database, Redis, auth, payment, email, business, and CORS configuration.",
    "backend/app/core/database.py": "Creates the async SQLAlchemy engine, session factory, and Base model class. It also switches pool behavior by environment and exposes get_db plus init_db helpers.",
    "backend/app/core/encryption.py": "Implements AES-256-GCM helpers and the EncryptedText SQLAlchemy type used for PII fields such as phone numbers and address lines.",
    "backend/app/core/openapi_config.py": "Contains the custom OpenAPI builder used to refine generated API documentation when that customization is enabled.",
    "backend/app/core/payment_config.py": "Defines payment-specific settings such as gateway credentials, enabled methods, retry windows, reservation timing, and currency defaults used by the payment layer.",
    "backend/app/core/security.py": "Provides password hashing, JWT creation and decoding, refresh-token hashing, and the low-level security helpers shared by auth flows.",
    "backend/app/middleware/auth.py": "Supplies dependency helpers for authenticated and role-based access. Endpoints use these functions to resolve the current user and enforce admin-only routes.",
    "backend/app/middleware/audit_log.py": "ASGI middleware that records admin actions into the audit trail. It extracts target metadata from requests and persists activity logs for privileged operations.",
    "backend/app/middleware/rate_limiter.py": "Redis-backed ASGI rate limiter for auth and general API traffic. It exempts selected session endpoints and adds limit headers to responses.",
    "backend/app/models/models.py": "Single-file ORM schema for the backend. It defines the full commerce, privacy, payment, audit, and settings data model used by SQLAlchemy and Alembic.",
    "backend/app/repositories/payment_repository.py": "Encapsulates payment-related database reads and writes so PaymentService can reuse consistent query and update primitives.",
    "backend/app/jobs/reservation_expiry.py": "Background stock-reservation cleanup job. It finds expired held reservations, restores inventory, and can run either periodically in-process or as a standalone task.",
    "backend/app/jobs/fx_rate_sync.py": "Background exchange-rate sync job. It refreshes cached FX rates and serves as the scheduled entrypoint for the FX service.",
    "backend/app/jobs/deletion_job.py": "Background account-deletion processor. It advances pending deletion requests past the grace period by calling the privacy-service anonymization workflow.",
    "backend/app/services/admin_dashboard_service.py": "Analytics and admin-operations service behind dashboard KPIs, reports, CSV exports, audit log search, and user management actions.",
    "backend/app/services/auth_service.py": "Core authentication service for registration, login, refresh-token rotation, logout, and phone or token guardrails around session creation.",
    "backend/app/services/cart_coupon_service.py": "Commerce service layer for server-side cart state and coupon lifecycle management. It handles guest-cart merge, quantity updates, coupon CRUD, and coupon application rules.",
    "backend/app/services/catalog_service.py": "Read-oriented catalog service for public storefront pages. It builds landing data, category browsing, filtered product lists, detail views, stock checks, and filter metadata.",
    "backend/app/services/email_verification_service.py": "Creates and validates email-verification records, including resend and token-based verification flows used during onboarding.",
    "backend/app/services/fx_rate_service.py": "Fetches, caches, and locks FX rates for non-INR checkout flows. Used by payment logic and the scheduled FX sync job.",
    "backend/app/services/image_service.py": "Admin image-pipeline service. It creates upload intents, records processing callbacks, manages ordering, and keeps primary-image state consistent.",
    "backend/app/services/invoice_sequence_service.py": "Small helper module for financial-year invoice numbering. It generates the next invoice number atomically and exposes current financial-year logic.",
    "backend/app/services/invoice_service.py": "Invoice and credit-note business logic. It assembles tax documents, numbering, PDF generation, listing, GST summaries, and regeneration workflows.",
    "backend/app/services/notification_service.py": "Notification façade for order lifecycle emails and future SMS hooks. It keeps transactional messaging separate from the core order and refund flows.",
    "backend/app/services/order_service.py": "Order lifecycle service covering checkout summaries, tax calculation, order placement, state transitions, timelines, and admin order retrieval.",
    "backend/app/services/order_state_machine.py": "Defines legal order-status transitions and helper checks such as allowed moves, terminal states, and cancellation rules.",
    "backend/app/services/password_reset_service.py": "Issues password-reset tokens and completes password resets after token validation.",
    "backend/app/services/payment_service.py": "Main payment orchestration service. It selects gateways, creates Razorpay or Stripe payment objects, supports UPI flows, processes webhooks, and drives refunds.",
    "backend/app/services/privacy_service.py": "Privacy and compliance workflow service. It manages consent state, account-deletion requests, anonymization, cookie preferences, and user data export.",
    "backend/app/services/product_service.py": "Admin product-management service for attribute definitions, categories, products, variants, inventory, size guides, bulk upload, and duplication workflows.",
    "backend/app/services/return_service.py": "Return and refund service layer. It validates return requests, performs approval or rejection steps, restocks inventory, and initiates gateway refunds plus credit notes.",
    "backend/app/services/shipping_service.py": "Shipment-management service handling shipment creation, status updates, delivery transitions, and courier integration stubs.",
    "backend/app/services/store_settings_service.py": "Runtime settings service for merchant UPI VPA, shipping configuration, seller information, contact details, and settings audit history.",
    "backend/app/services/totp_service.py": "Admin two-factor authentication service. It handles TOTP setup, validation, disablement, and 2FA-required checks by role.",
    "backend/app/services/wishlist_review_service.py": "Review moderation and customer review service. It powers review listing, creation, editing, deletion, and admin moderation queues.",
    "backend/app/services/gateways/razorpay_client.py": "Thin provider wrapper around the Razorpay SDK and REST calls used for India checkout, UPI polling, webhook verification, refunds, and QR generation.",
    "backend/app/services/gateways/stripe_client.py": "Thin provider wrapper around Stripe initialization, PaymentIntent creation, webhook verification, and refund operations for global payments.",
    "backend/app/api/v1/endpoints/auth.py": "FastAPI auth route module for registration, login, refresh, logout, current-user lookup, email verification, and password-reset flows. It also sets and clears the httpOnly auth cookies.",
    "backend/app/api/v1/endpoints/totp.py": "FastAPI admin 2FA route module for TOTP setup, verification, login validation, disablement, and status lookup.",
    "backend/app/api/v1/endpoints/user.py": "FastAPI profile route module for reading and updating the logged-in user profile, changing passwords, and managing customer addresses.",
    "backend/app/api/v1/endpoints/privacy.py": "FastAPI compliance route module for consent retrieval or updates, deletion requests, deletion status, data export, and cookie-consent persistence.",
    "backend/app/api/v1/endpoints/catalog.py": "Public storefront route module for landing data, category browsing, product lists and detail views, search autocomplete, size guides, stock checks, and filter metadata.",
    "backend/app/api/v1/endpoints/wishlist_reviews.py": "Customer and admin review route module. It exposes review creation and moderation endpoints used around product-review workflows.",
    "backend/app/api/v1/endpoints/cart_coupons.py": "Cart and coupon route module for cart CRUD, guest-cart merge, coupon application, and admin coupon maintenance operations.",
    "backend/app/api/v1/endpoints/orders.py": "Checkout and order-management route module. It covers address CRUD, order summary, place-order, customer order history, cancellation, and admin transition endpoints.",
    "backend/app/api/v1/endpoints/payments.py": "Payment route module exposing gateway selection, Razorpay and Stripe creation flows, UPI collect or QR polling, status checks, FX locks, webhooks, and admin refund initiation.",
    "backend/app/api/v1/endpoints/invoices.py": "Invoice and credit-note route module for customer downloads, admin searches, presigned access, GST summaries, and regeneration actions.",
    "backend/app/api/v1/endpoints/shipping_returns.py": "Shipping, return, and refund route module for customer return requests plus admin shipment, return, and refund operations.",
    "backend/app/api/v1/endpoints/admin_images.py": "Admin image-pipeline route module for presigned uploads, processing callbacks, image metadata edits, reordering, deletion, and primary-image selection.",
    "backend/app/api/v1/endpoints/admin_products.py": "Largest admin-commerce route module. It exposes attribute, category, product, variant, inventory, size-guide, bulk-upload, bulk-create, and duplication endpoints.",
    "backend/app/api/v1/endpoints/admin_dashboard.py": "Admin reporting route module for dashboard stats, revenue trends, top products, low-stock alerts, reports, exports, audit logs, and user-management actions.",
    "backend/app/api/v1/endpoints/admin_settings.py": "Admin configuration route module for merchant UPI, shipping, seller, and contact settings plus audit-trail reads.",
    "backend/app/utils/soft_delete.py": "Utility helpers for the repository-wide soft-delete pattern. It provides the mixin and event hooks that filter deleted rows by default.",
    "backend/app/utils/seed.py": "Seed script for permissions and invoice-sequence bootstrap data. It is intended for one-time or controlled environment initialization.",
    "backend/app/utils/seed_attributes.py": "Seed script for default apparel attribute definitions so category and filter metadata can be initialized quickly.",
    "backend/app/utils/pdf_generator.py": "Renders invoice and credit-note PDFs and uploads them to S3. It also keeps HTML and PDF generation helpers together for tax-document output.",
    "backend/app/utils/number_to_words.py": "Converts monetary amounts into Indian rupee wording for invoice-style document output.",
    "backend/app/utils/gst_utils.py": "GST helper module for state-code lookup, supply-type detection, and intra- or inter-state tax component splitting.",
    "backend/app/utils/encryption.py": "Compatibility wrapper that preserves an older import path while delegating real encryption behavior to app.core.encryption.",
    "backend/app/utils/email_sender.py": "Transactional email helper for OTP, verification, password-reset, and general message sends through the configured email backend.",
    "backend/app/tests/test_payment_service.py": "Pytest coverage for the payment service, including gateway selection, order creation, webhook idempotency, refunds, and FX-related behavior.",
    "backend/app/tests/test_pii_encryption.py": "Pytest coverage for the PII encryption helpers and EncryptedText behavior, including query-time comparisons and invalid ciphertext handling.",
    "backend/aws/lambda/image-processor/lambda_function.py": "Standalone AWS Lambda processor for product-image post-processing. It performs the image task and posts results back to the backend callback endpoint.",
    "backend/migrations/env.py": "Alembic environment configuration. It wires migration metadata and controls online or offline migration execution.",
    "backend/migrations/versions/44600a47af5d_initial_schema.py": "Initial Alembic revision that creates the baseline application schema and related indexes or constraints.",
    "backend/migrations/versions/b8d9a41f2c6e_expand_pii_columns_to_text.py": "Follow-up Alembic revision that widens selected encrypted PII columns from fixed strings to text so ciphertext fits safely.",
}


SECTIONS = [
    (
        "Application Bootstrap",
        [
            "backend/app/__init__.py",
            "backend/app/api/__init__.py",
            "backend/app/api/v1/__init__.py",
            "backend/app/api/v1/router.py",
            "backend/app/main.py",
            "backend/lambda_handler.py",
        ],
        "These files create the application entrypoints and package structure used by local FastAPI startup and AWS Lambda execution.",
    ),
    (
        "Core",
        [
            "backend/app/core/__init__.py",
            "backend/app/core/config.py",
            "backend/app/core/database.py",
            "backend/app/core/encryption.py",
            "backend/app/core/openapi_config.py",
            "backend/app/core/payment_config.py",
            "backend/app/core/security.py",
        ],
        "The core layer provides shared configuration, database connectivity, encryption, security, and API metadata primitives.",
    ),
    (
        "Middleware",
        [
            "backend/app/middleware/__init__.py",
            "backend/app/middleware/auth.py",
            "backend/app/middleware/audit_log.py",
            "backend/app/middleware/rate_limiter.py",
        ],
        "Middleware and auth-dependency modules apply cross-cutting concerns such as access control, audit logging, and rate limiting.",
    ),
    (
        "Models and Repository",
        [
            "backend/app/models/__init__.py",
            "backend/app/models/models.py",
            "backend/app/repositories/__init__.py",
            "backend/app/repositories/payment_repository.py",
        ],
        "This layer defines the SQLAlchemy schema and the repository helper used by the payment flow.",
    ),
    (
        "Schemas",
        [
            "backend/app/schemas/__init__.py",
            "backend/app/schemas/admin_dashboard_schemas.py",
            "backend/app/schemas/auth.py",
            "backend/app/schemas/cart_coupon.py",
            "backend/app/schemas/catalog.py",
            "backend/app/schemas/image.py",
            "backend/app/schemas/invoice_schemas.py",
            "backend/app/schemas/order.py",
            "backend/app/schemas/payment.py",
            "backend/app/schemas/privacy_schemas.py",
            "backend/app/schemas/product.py",
            "backend/app/schemas/shipping_returns_schemas.py",
            "backend/app/schemas/wishlist_review.py",
        ],
        "Schema modules define the Pydantic request and response contracts used by the endpoint layer.",
    ),
    (
        "Services",
        [
            "backend/app/services/__init__.py",
            "backend/app/services/admin_dashboard_service.py",
            "backend/app/services/auth_service.py",
            "backend/app/services/cart_coupon_service.py",
            "backend/app/services/catalog_service.py",
            "backend/app/services/email_verification_service.py",
            "backend/app/services/fx_rate_service.py",
            "backend/app/services/image_service.py",
            "backend/app/services/invoice_sequence_service.py",
            "backend/app/services/invoice_service.py",
            "backend/app/services/notification_service.py",
            "backend/app/services/order_service.py",
            "backend/app/services/order_state_machine.py",
            "backend/app/services/password_reset_service.py",
            "backend/app/services/payment_service.py",
            "backend/app/services/privacy_service.py",
            "backend/app/services/product_service.py",
            "backend/app/services/return_service.py",
            "backend/app/services/shipping_service.py",
            "backend/app/services/store_settings_service.py",
            "backend/app/services/totp_service.py",
            "backend/app/services/wishlist_review_service.py",
        ],
        "Service modules hold the business logic behind each commerce, auth, privacy, reporting, and fulfillment workflow.",
    ),
    (
        "Gateway Clients",
        [
            "backend/app/services/gateways/__init__.py",
            "backend/app/services/gateways/razorpay_client.py",
            "backend/app/services/gateways/stripe_client.py",
        ],
        "These modules isolate provider-specific payment SDK calls from the rest of the application.",
    ),
    (
        "API Endpoints",
        [
            "backend/app/api/v1/endpoints/__init__.py",
            "backend/app/api/v1/endpoints/admin_dashboard.py",
            "backend/app/api/v1/endpoints/admin_images.py",
            "backend/app/api/v1/endpoints/admin_products.py",
            "backend/app/api/v1/endpoints/admin_settings.py",
            "backend/app/api/v1/endpoints/auth.py",
            "backend/app/api/v1/endpoints/cart_coupons.py",
            "backend/app/api/v1/endpoints/catalog.py",
            "backend/app/api/v1/endpoints/invoices.py",
            "backend/app/api/v1/endpoints/orders.py",
            "backend/app/api/v1/endpoints/payments.py",
            "backend/app/api/v1/endpoints/privacy.py",
            "backend/app/api/v1/endpoints/shipping_returns.py",
            "backend/app/api/v1/endpoints/totp.py",
            "backend/app/api/v1/endpoints/user.py",
            "backend/app/api/v1/endpoints/wishlist_reviews.py",
        ],
        "Endpoint modules are the public HTTP surface of the backend. They validate requests, enforce auth, and delegate to service modules.",
    ),
    (
        "Jobs",
        [
            "backend/app/jobs/__init__.py",
            "backend/app/jobs/deletion_job.py",
            "backend/app/jobs/fx_rate_sync.py",
            "backend/app/jobs/reservation_expiry.py",
        ],
        "Background job modules support periodic cleanup and scheduled processing outside the request-response path.",
    ),
    (
        "Utilities",
        [
            "backend/app/utils/__init__.py",
            "backend/app/utils/email_sender.py",
            "backend/app/utils/encryption.py",
            "backend/app/utils/gst_utils.py",
            "backend/app/utils/number_to_words.py",
            "backend/app/utils/pdf_generator.py",
            "backend/app/utils/seed.py",
            "backend/app/utils/seed_attributes.py",
            "backend/app/utils/soft_delete.py",
        ],
        "Utility modules contain shared helpers, document generation, and environment bootstrapping scripts that are reused across services.",
    ),
    (
        "AWS Auxiliary Lambda",
        [
            "backend/aws/lambda/image-processor/lambda_function.py",
        ],
        "The repo also ships a dedicated Lambda for image post-processing outside the main API container.",
    ),
    (
        "Tests",
        [
            "backend/app/tests/test_payment_service.py",
            "backend/app/tests/test_pii_encryption.py",
        ],
        "The current checked-in Python test suite is small and focused on the payment service and PII encryption behavior.",
    ),
    (
        "Migrations",
        [
            "backend/migrations/env.py",
            "backend/migrations/versions/44600a47af5d_initial_schema.py",
            "backend/migrations/versions/b8d9a41f2c6e_expand_pii_columns_to_text.py",
        ],
        "Alembic files define schema-migration wiring and the concrete revisions checked into the repository.",
    ),
]


class FileMeta:
    def __init__(self, path: str, doc: str, classes: list[str], functions: list[str],
                 async_functions: list[str], route_functions: list[str],
                 class_methods: dict[str, list[str]]) -> None:
        self.path = path
        self.doc = doc
        self.classes = classes
        self.functions = functions
        self.async_functions = async_functions
        self.route_functions = route_functions
        self.class_methods = class_methods


def set_run_font(run, name: str = "Calibri", size: int | None = None,
                 color: RGBColor | None = None, bold: bool | None = None,
                 italic: bool | None = None) -> None:
    run.font.name = name
    run._element.rPr.rFonts.set(qn("w:ascii"), name)
    run._element.rPr.rFonts.set(qn("w:hAnsi"), name)
    if size is not None:
        run.font.size = Pt(size)
    if color is not None:
        run.font.color.rgb = color
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic


def shade_cell(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def set_cell_margins(cell, top=80, start=120, bottom=80, end=120) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for key, value in {"top": top, "start": start, "bottom": bottom, "end": end}.items():
        node = tc_mar.find(qn(f"w:{key}"))
        if node is None:
            node = OxmlElement(f"w:{key}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def format_cell(cell, fill: str | None = None, bold: bool = False,
                font_size: int = 9, color: RGBColor = DARK) -> None:
    if fill:
        shade_cell(cell, fill)
    set_cell_margins(cell)
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    for paragraph in cell.paragraphs:
        paragraph.paragraph_format.space_before = Pt(0)
        paragraph.paragraph_format.space_after = Pt(0)
        paragraph.paragraph_format.line_spacing = 1.05
        for run in paragraph.runs:
            set_run_font(run, size=font_size, color=color, bold=bold)


def set_repeat_table_header(row) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)


def set_table_layout(table, widths: list[float]) -> None:
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    for row in table.rows:
        for index, width in enumerate(widths):
            row.cells[index].width = Inches(width)


def add_title(doc: Document, text: str, size: int = 24, color: RGBColor = NAVY,
              after: int = 6) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(after)
    run = p.add_run(text)
    set_run_font(run, size=size, color=color, bold=True)


def add_paragraph(doc: Document, text: str, size: int = 11, color: RGBColor = DARK,
                  bold: bool = False, italic: bool = False, after: int = 6) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(after)
    p.paragraph_format.line_spacing = 1.1
    run = p.add_run(text)
    set_run_font(run, size=size, color=color, bold=bold, italic=italic)


def add_callout(doc: Document, title: str, body: str) -> None:
    table = doc.add_table(rows=2, cols=1)
    set_table_layout(table, [6.5])
    table.style = "Table Grid"
    header = table.cell(0, 0)
    header.text = title
    format_cell(header, fill=LIGHT_FILL, bold=True, font_size=11, color=NAVY)
    body_cell = table.cell(1, 0)
    body_cell.text = body
    format_cell(body_cell, fill=CALLOUT_FILL, font_size=10, color=DARK)
    doc.add_paragraph().paragraph_format.space_after = Pt(4)


def configure_document(doc: Document) -> None:
    section = doc.sections[0]
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)

    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(11)

    for style_name, size, color in [
        ("Heading 1", 16, BLUE),
        ("Heading 2", 13, BLUE),
        ("Heading 3", 11, NAVY),
    ]:
        style = doc.styles[style_name]
        style.font.name = "Calibri"
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = color
        style.paragraph_format.space_before = Pt(12 if style_name != "Heading 3" else 8)
        style.paragraph_format.space_after = Pt(6 if style_name != "Heading 3" else 4)

    footer = section.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = footer.add_run(
        f"Ashmiwebportal Python File Reference | Generated from current repo snapshot | {DATE_LABEL}"
    )
    set_run_font(run, size=9, color=MUTED)


def parse_python_file(path: Path) -> FileMeta:
    text = path.read_text(encoding="utf-8-sig", errors="ignore")
    tree = ast.parse(text)
    doc = ast.get_docstring(tree) or ""
    classes: list[str] = []
    functions: list[str] = []
    async_functions: list[str] = []
    route_functions: list[str] = []
    class_methods: dict[str, list[str]] = {}

    for node in tree.body:
        if isinstance(node, ast.ClassDef):
            classes.append(node.name)
            methods = [
                child.name
                for child in node.body
                if isinstance(child, (ast.FunctionDef, ast.AsyncFunctionDef)) and not child.name.startswith("__")
            ]
            if methods:
                class_methods[node.name] = methods
        elif isinstance(node, ast.FunctionDef):
            functions.append(node.name)
            if _is_route_handler(node):
                route_functions.append(node.name)
        elif isinstance(node, ast.AsyncFunctionDef):
            async_functions.append(node.name)
            if _is_route_handler(node):
                route_functions.append(node.name)

    return FileMeta(
        path=str(path.relative_to(ROOT)).replace("\\", "/"),
        doc=doc.splitlines()[0] if doc else "",
        classes=classes,
        functions=functions,
        async_functions=async_functions,
        route_functions=route_functions,
        class_methods=class_methods,
    )


def _is_route_handler(node: ast.FunctionDef | ast.AsyncFunctionDef) -> bool:
    for dec in node.decorator_list:
        target = dec.func if isinstance(dec, ast.Call) else dec
        if isinstance(target, ast.Attribute) and target.attr in {"get", "post", "put", "patch", "delete"}:
            return True
    return False


def domain_from_stem(stem: str) -> str:
    text = stem.replace("_", " ")
    replacements = {
        "totp": "TOTP and admin 2FA",
        "auth": "authentication",
        "catalog": "public catalog",
        "cart coupons": "cart and coupons",
        "wishlist reviews": "reviews",
        "shipping returns": "shipping and returns",
        "admin dashboard": "admin dashboards and reports",
        "admin products": "admin products and inventory",
        "admin settings": "admin store settings",
        "admin images": "admin image processing",
        "order": "orders",
        "payment": "payments",
        "product": "products",
    }
    return replacements.get(text, text)


def summarize(meta: FileMeta) -> str:
    if meta.path in OVERRIDES:
        return OVERRIDES[meta.path]

    if meta.path.endswith("__init__.py"):
        parent = Path(meta.path).parent.as_posix()
        return f"Package initializer for `{parent}`. It mainly keeps the folder importable and does not contain substantive runtime logic."

    if "/schemas/" in meta.path:
        domain = domain_from_stem(Path(meta.path).stem)
        return f"Defines Pydantic request and response models for {domain}. Endpoint modules use these classes for validation, serialization, and typed API contracts."

    if "/api/v1/endpoints/" in meta.path:
        domain = domain_from_stem(Path(meta.path).stem)
        count = len(meta.route_functions)
        return f"FastAPI route module for {domain}. It mounts {count} HTTP handler{'s' if count != 1 else ''} under the versioned API and delegates business logic to services."

    if "/services/gateways/" in meta.path:
        domain = domain_from_stem(Path(meta.path).stem)
        return f"Gateway-client module for {domain}. It isolates direct provider calls so the payment orchestration layer stays provider-agnostic."

    if "/services/" in meta.path:
        domain = domain_from_stem(Path(meta.path).stem)
        return f"Business-logic service module for {domain}. It is consumed by the endpoint layer and works with the async database session or external helpers."

    if "/jobs/" in meta.path:
        return meta.doc or "Background-job module used by scheduled or periodic processing."

    if "/tests/" in meta.path:
        domain = domain_from_stem(Path(meta.path).stem.replace("test_", ""))
        return f"Pytest module that verifies {domain}-related behavior in the backend."

    if "/migrations/versions/" in meta.path:
        return f"Alembic revision script used to evolve the database schema. This file contains the upgrade and downgrade steps for its revision."

    if meta.path.endswith("migrations/env.py"):
        return "Alembic environment bootstrap that connects migration execution to the application's model metadata."

    if "/utils/" in meta.path:
        domain = domain_from_stem(Path(meta.path).stem)
        return f"Utility module for {domain}. It provides shared helper logic that is reused by services, jobs, or scripts."

    return meta.doc or "Python module in the backend codebase."


def summarize_symbols(meta: FileMeta) -> str:
    if meta.route_functions:
        return "Handlers: " + limit_items(meta.route_functions, 6)

    if meta.class_methods:
        parts = []
        for cls, methods in meta.class_methods.items():
            parts.append(f"{cls}: {limit_items(methods, 5)}")
        return "Key methods: " + " | ".join(parts[:2])

    if meta.classes:
        return "Classes: " + limit_items(meta.classes, 6)

    names = meta.functions + meta.async_functions
    if names:
        return "Functions: " + limit_items(names, 6)

    return "No substantive top-level symbols; package or marker module."


def limit_items(items: list[str], limit: int) -> str:
    shown = items[:limit]
    extra = len(items) - len(shown)
    text = ", ".join(shown)
    if extra > 0:
        text += f" (+{extra} more)"
    return text


def collect_metadata() -> dict[str, FileMeta]:
    result: dict[str, FileMeta] = {}
    for path in sorted(ROOT.glob("backend/**/*.py")):
        if ".venv" in path.parts or "__pycache__" in path.parts:
            continue
        meta = parse_python_file(path)
        result[meta.path] = meta
    return result


def add_section_table(doc: Document, files: list[str], meta_map: dict[str, FileMeta]) -> None:
    table = doc.add_table(rows=1, cols=2)
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    widths = [2.0, 4.5]
    set_table_layout(table, widths)

    headers = ["File", "Usage and key symbols"]
    for idx, text in enumerate(headers):
        cell = table.rows[0].cells[idx]
        cell.text = text
        format_cell(cell, fill=LIGHT_FILL, bold=True, font_size=10, color=NAVY)
    set_repeat_table_header(table.rows[0])

    for rel in files:
        meta = meta_map[rel]
        row = table.add_row().cells
        row[0].text = rel
        file_p = row[0].paragraphs[0]
        file_p.paragraph_format.space_after = Pt(0)
        file_p.paragraph_format.line_spacing = 1.0
        for run in file_p.runs:
            set_run_font(run, size=8.5, color=NAVY, bold=True)
        format_cell(row[0], font_size=8.5, color=NAVY)

        usage = summarize(meta)
        symbols = summarize_symbols(meta)

        p1 = row[1].paragraphs[0]
        p1.text = ""
        run1 = p1.add_run(usage)
        set_run_font(run1, size=9, color=DARK)
        p1.paragraph_format.space_after = Pt(2)
        p1.paragraph_format.line_spacing = 1.05

        p2 = row[1].add_paragraph()
        run2 = p2.add_run(symbols)
        set_run_font(run2, size=8.5, color=MUTED, italic=True)
        p2.paragraph_format.space_before = Pt(0)
        p2.paragraph_format.space_after = Pt(0)
        p2.paragraph_format.line_spacing = 1.0
        format_cell(row[1], font_size=9, color=DARK)

    doc.add_paragraph().paragraph_format.space_after = Pt(4)


def build_document(out_path: Path) -> None:
    meta_map = collect_metadata()
    doc = Document()
    configure_document(doc)

    total_files = len(meta_map)
    add_paragraph(doc, "Detailed Python Reference", size=12, color=BLUE, bold=True, after=2)
    add_title(doc, "Ashmiwebportal", size=26, color=NAVY, after=4)
    add_paragraph(
        doc,
        "Backend Python file-by-file usage guide generated from the current repository contents.",
        size=13,
        color=MUTED,
        after=4,
    )
    add_paragraph(
        doc,
        f"Snapshot date: {DATE_LABEL}. Scope: {total_files} checked-in `.py` files under `backend/`. The `frontend/` application currently contains no Python modules.",
        size=10,
        color=MUTED,
        italic=True,
        after=12,
    )

    add_callout(
        doc,
        "What changed from the first draft",
        "This edition expands the earlier high-level summary into a full Python module reference. "
        "Every checked-in backend `.py` file is now listed with its usage and its main exported symbols or handlers.",
    )

    doc.add_heading("Reading Guide", level=1)
    add_paragraph(
        doc,
        "The tables below are grouped by layer so the document reads like the original codebase summary, but at file granularity. "
        "Each row explains what the file is for and then calls out the most important classes, functions, methods, or route handlers defined inside it.",
        after=8,
    )

    for idx, (title, files, blurb) in enumerate(SECTIONS, start=1):
        doc.add_heading(f"{idx}. {title}", level=1)
        add_paragraph(doc, blurb, after=6)
        add_section_table(doc, files, meta_map)

    doc.add_heading(f"{len(SECTIONS) + 1}. Closing Notes", level=1)
    add_paragraph(
        doc,
        "This reference is intentionally Python-only because that was the missing level of detail requested. "
        "If you want, the next iteration can add the same file-by-file treatment for the React frontend, including every `.jsx`, `.js`, and config file.",
        after=0,
    )

    doc.save(out_path)


def main() -> None:
    output_dir = ROOT / "generated"
    output_dir.mkdir(exist_ok=True)
    out_path = output_dir / "Ashmiwebportal_Python_File_Reference_2026-05-14.docx"
    build_document(out_path)
    print(out_path)


if __name__ == "__main__":
    main()
