from __future__ import annotations

import json
import re
from datetime import datetime
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION_START
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path.cwd()
DATE_LABEL = "14 May 2026"

BLUE = RGBColor(46, 116, 181)
NAVY = RGBColor(11, 37, 69)
DARK = RGBColor(34, 34, 34)
MUTED = RGBColor(90, 90, 90)
LIGHT_FILL = "F2F4F7"
CALLOUT_FILL = "F7F9FC"


def read_text(path: Path) -> str:
    return path.read_text(encoding="utf-8", errors="ignore")


def set_run_font(run, name: str = "Calibri", size: int | None = None,
                 color: RGBColor | None = None, bold: bool | None = None,
                 italic: bool | None = None) -> None:
    run.font.name = name
    run._element.rPr.rFonts.set(qn("w:ascii"), name)
    run._element.rPr.rFonts.set(qn("w:hAnsi"), name)
    if size is not None:
        run.font.size = Pt(size)
    if color is not None:
        run.font.color.rgb = color
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic


def shade_cell(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def set_cell_margins(cell, top=80, start=120, bottom=80, end=120) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for key, value in {"top": top, "start": start, "bottom": bottom, "end": end}.items():
        node = tc_mar.find(qn(f"w:{key}"))
        if node is None:
            node = OxmlElement(f"w:{key}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def format_cell(cell, fill: str | None = None, bold: bool = False,
                font_size: int = 10, color: RGBColor = DARK) -> None:
    if fill:
        shade_cell(cell, fill)
    set_cell_margins(cell)
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    for paragraph in cell.paragraphs:
        paragraph.paragraph_format.space_before = Pt(0)
        paragraph.paragraph_format.space_after = Pt(0)
        paragraph.paragraph_format.line_spacing = 1.1
        for run in paragraph.runs:
            set_run_font(run, size=font_size, color=color, bold=bold)


def set_table_layout(table, widths_in_inches: list[float]) -> None:
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    for row in table.rows:
        for index, width in enumerate(widths_in_inches):
            row.cells[index].width = Inches(width)


def set_repeat_table_header(row) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)


def add_title(doc: Document, text: str, size: int = 24, color: RGBColor = NAVY,
              after: int = 6, center: bool = False) -> None:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER if center else WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(after)
    run = p.add_run(text)
    set_run_font(run, size=size, color=color, bold=True)


def add_paragraph(doc: Document, text: str, size: int = 11, color: RGBColor = DARK,
                  bold: bool = False, italic: bool = False, after: int = 6,
                  center: bool = False) -> None:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER if center else WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(after)
    p.paragraph_format.line_spacing = 1.1
    run = p.add_run(text)
    set_run_font(run, size=size, color=color, bold=bold, italic=italic)


def add_bullets(doc: Document, items: list[str]) -> None:
    for item in items:
        p = doc.add_paragraph(style="List Bullet")
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(4)
        p.paragraph_format.line_spacing = 1.15
        run = p.add_run(item)
        set_run_font(run, size=11, color=DARK)


def add_callout(doc: Document, title: str, body: str) -> None:
    table = doc.add_table(rows=2, cols=1)
    set_table_layout(table, [6.5])
    table.style = "Table Grid"
    header = table.cell(0, 0)
    header.text = title
    format_cell(header, fill=LIGHT_FILL, bold=True, font_size=11, color=NAVY)
    body_cell = table.cell(1, 0)
    body_cell.text = body
    format_cell(body_cell, fill=CALLOUT_FILL, font_size=10, color=DARK)
    doc.add_paragraph().paragraph_format.space_after = Pt(4)


def add_table(doc: Document, headers: list[str], rows: list[list[str]],
              widths: list[float], font_size: int = 10) -> None:
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    set_table_layout(table, widths)
    header_cells = table.rows[0].cells
    for idx, text in enumerate(headers):
        header_cells[idx].text = text
        format_cell(header_cells[idx], fill=LIGHT_FILL, bold=True, font_size=font_size, color=NAVY)
    set_repeat_table_header(table.rows[0])

    for row_values in rows:
        row = table.add_row().cells
        for idx, text in enumerate(row_values):
            row[idx].text = text
            format_cell(row[idx], font_size=font_size, color=DARK)

    doc.add_paragraph().paragraph_format.space_after = Pt(4)


def configure_document(doc: Document) -> None:
    section = doc.sections[0]
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)

    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(11)

    for style_name, size, color in [
        ("Heading 1", 16, BLUE),
        ("Heading 2", 13, BLUE),
        ("Heading 3", 12, NAVY),
    ]:
        style = doc.styles[style_name]
        style.font.name = "Calibri"
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = color
        style.paragraph_format.space_before = Pt(12 if style_name != "Heading 3" else 8)
        style.paragraph_format.space_after = Pt(6 if style_name != "Heading 3" else 4)

    footer = section.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    footer_run = footer.add_run(
        f"Ashmiwebportal Codebase Summary | Generated from current repo snapshot | {DATE_LABEL}"
    )
    set_run_font(footer_run, size=9, color=MUTED)


def parse_route_count() -> int:
    count = 0
    endpoint_dir = ROOT / "backend" / "app" / "api" / "v1" / "endpoints"
    pattern = re.compile(r"@\w+[\w_]*\.(get|post|put|patch|delete)\(")
    for file in endpoint_dir.glob("*.py"):
        if file.name == "__init__.py":
            continue
        count += len(pattern.findall(read_text(file)))
    return count


def parse_models() -> list[str]:
    text = read_text(ROOT / "backend" / "app" / "models" / "models.py")
    return re.findall(r"^class\s+(\w+)\(", text, flags=re.MULTILINE)


def collect_metrics() -> dict:
    frontend_package = json.loads(read_text(ROOT / "frontend" / "package.json"))
    workflow = read_text(ROOT / ".github" / "workflows" / "deploy-dev-v3.yml")
    requirements = read_text(ROOT / "backend" / "requirements.txt")

    route_count = parse_route_count()
    models = parse_models()
    service_files = sorted(
        p.name for p in (ROOT / "backend" / "app" / "services").glob("*.py") if p.name != "__init__.py"
    )
    endpoint_files = sorted(
        p.name for p in (ROOT / "backend" / "app" / "api" / "v1" / "endpoints").glob("*.py") if p.name != "__init__.py"
    )
    page_files = sorted((ROOT / "frontend" / "src" / "pages").rglob("*.jsx"))
    store_files = sorted(p.name for p in (ROOT / "frontend" / "src" / "stores").glob("*.js"))
    e2e_files = sorted(p.name for p in (ROOT / "frontend" / "src" / "tests" / "e2e").glob("*"))
    unit_files = sorted(p.name for p in (ROOT / "frontend" / "src" / "tests" / "unit").glob("*"))

    workflow_branch = "develop" if "branches:\n      - develop" in workflow else "not detected"

    return {
        "generated_on": DATE_LABEL,
        "python": "Python 3.11",
        "backend_framework": f"FastAPI {requirements_match(requirements, 'fastapi')}",
        "frontend_framework": f"React {version_major(frontend_package['dependencies'].get('react', ''))}",
        "vite_version": frontend_package["devDependencies"].get("vite", ""),
        "tailwind_version": frontend_package["devDependencies"].get("tailwindcss", ""),
        "route_count": route_count,
        "model_count": len(models),
        "models": models,
        "endpoint_files": endpoint_files,
        "service_files": service_files,
        "page_count": len(page_files),
        "page_files": page_files,
        "store_files": store_files,
        "e2e_count": len(e2e_files),
        "unit_count": len(unit_files),
        "workflow_branch": workflow_branch,
    }


def requirements_match(requirements: str, package: str) -> str:
    match = re.search(rf"^{re.escape(package)}==([^\r\n]+)", requirements, flags=re.MULTILINE)
    return match.group(1) if match else "unknown"


def version_major(version: str) -> str:
    return version.lstrip("^~")


def build_doc(metrics: dict, out_path: Path) -> None:
    doc = Document()
    configure_document(doc)

    add_paragraph(doc, "Functional Summary", size=12, color=BLUE, bold=True, after=2)
    add_title(doc, "Ashmiwebportal", size=26, color=NAVY, after=4)
    add_paragraph(
        doc,
        "Current codebase summary generated from the checked-in project state under C:\\Ashmiwebportal",
        size=13,
        color=MUTED,
        after=2,
    )
    add_paragraph(
        doc,
        "Reference PDF superseded: March 2026 summary. This edition reflects the repository snapshot dated 14 May 2026.",
        size=10,
        color=MUTED,
        italic=True,
        after=14,
    )

    add_table(
        doc,
        ["Area", "Current snapshot"],
        [
            ["Backend", "FastAPI 0.115.6, SQLAlchemy async, Redis rate limiting, Lambda entrypoint via Mangum"],
            ["Frontend", "React 19.2, Vite 7.3, Tailwind CSS v4, Zustand stores, React Router"],
            ["Persistence", "PostgreSQL via asyncpg at runtime, psycopg2 for Alembic, 33 ORM table classes in models.py"],
            ["Payments", "Razorpay for India and UPI flows, Stripe for non-India card flows, webhook idempotency via PaymentEvent"],
            ["Ops surface", "Docker Compose for local API stack, GitHub Actions deploy workflow, S3 + CloudFront frontend publish"],
        ],
        [1.7, 4.8],
    )

    add_callout(
        doc,
        "Update note",
        "The attached March 2026 summary no longer matches the checked-in repository in several places. "
        "This updated document is based on the current source tree, key entrypoints, package manifests, "
        "workflow configuration, and measured file counts rather than prior narrative assumptions.",
    )

    doc.add_heading("1. Repository Snapshot", level=1)
    add_paragraph(
        doc,
        "The project remains a two-part application: a Python backend under backend/ and a React SPA under frontend/. "
        "Current top-level repo assets also include Docker Compose, GitHub Actions deployment workflow, handoff/runbook documentation, and a checked-in virtual environment.",
        after=8,
    )
    add_table(
        doc,
        ["Metric", "Value", "How measured"],
        [
            ["Explicit API handlers", str(metrics["route_count"]), "Route decorators counted in backend/app/api/v1/endpoints/*.py"],
            ["Endpoint modules", str(len(metrics["endpoint_files"])), "Python modules under backend/app/api/v1/endpoints excluding __init__.py"],
            ["ORM table classes", str(metrics["model_count"]), "Class declarations in backend/app/models/models.py"],
            ["Service modules", str(len(metrics["service_files"])), "Top-level service .py files under backend/app/services excluding __init__.py"],
            ["Frontend page components", str(metrics["page_count"]), "JSX files under frontend/src/pages"],
            ["Frontend stores", str(len(metrics["store_files"])), "Zustand store files under frontend/src/stores"],
            ["E2E specs", str(metrics["e2e_count"]), "Files under frontend/src/tests/e2e"],
            ["Unit test files", str(metrics["unit_count"]), "Files under frontend/src/tests/unit"],
        ],
        [2.0, 1.0, 3.5],
    )
    add_bullets(
        doc,
        [
            "Local backend entrypoint remains uvicorn app.main:app with Docker Compose services for Postgres and Redis.",
            "The Vite dev server is configured for port 3000 and proxies /api requests to http://localhost:8000.",
            "Playwright is configured to run from frontend/src/tests/e2e with Desktop Chrome and Mobile Chrome projects.",
            "The checked-in GitHub Actions deployment workflow currently triggers from the develop branch, not main.",
        ],
    )

    doc.add_heading("2. Backend Architecture", level=1)
    add_paragraph(
        doc,
        "The backend is organized as a conventional FastAPI service with distinct core, middleware, models, schemas, services, repositories, jobs, and API endpoint layers. "
        "app/main.py builds the FastAPI application, registers CORS, Redis-backed rate limiting, and admin audit logging, then mounts the aggregated v1 router plus dedicated payment routers.",
        after=8,
    )

    doc.add_heading("Core and runtime behavior", level=2)
    add_bullets(
        doc,
        [
            "backend/app/core/config.py centralizes environment-driven configuration, including DB URLs, JWT settings, payment keys, Redis URL, CORS origins, and encryption inputs.",
            "backend/app/core/database.py switches pool behavior by environment: AsyncAdaptedQueuePool for local development and NullPool with SSL for aws_dev or production Lambda execution.",
            "backend/lambda_handler.py uses lazy app loading with Mangum for HTTP requests and separately routes EventBridge-style task payloads for reservation release, FX sync, and deletion processing.",
            "backend/app/jobs/reservation_expiry.py runs inside the FastAPI lifespan loop every 60 seconds to release expired stock reservations while the app process is up.",
        ],
    )

    doc.add_heading("Domain coverage", level=2)
    add_bullets(
        doc,
        [
            "Authentication includes registration, login, refresh, password reset, email verification, and admin TOTP flows.",
            "Commerce flows cover catalog browsing, dynamic product attributes, cart and coupon handling, checkout, order placement, and order timeline retrieval.",
            "Payment orchestration supports Razorpay checkout, UPI collect or QR polling, Stripe intents, refund initiation, and idempotent webhook processing.",
            "Operational domains include invoice generation, shipment management, return handling, privacy workflows, dashboard reporting, and mutable store settings.",
        ],
    )

    doc.add_heading("Data model", level=2)
    add_paragraph(
        doc,
        "The single models.py file currently defines 33 ORM table classes. The schema spans customer identity, privacy and deletion workflow, catalog, pricing, cart, inventory reservation, order snapshots, invoices and credit notes, returns, refunds, payments, audit logging, and runtime store settings.",
        after=6,
    )
    add_table(
        doc,
        ["Domain group", "Representative tables"],
        [
            ["Identity and privacy", "User, UserAddress, UserConsent, AccountDeletionRequest, RefreshToken, EmailVerification"],
            ["Catalog and inventory", "Category, Product, AttributeDefinition, ProductVariant, ProductImage, SizeGuide, InventoryReservation"],
            ["Cart and promotions", "Cart, CartItem, Coupon, CouponUsage"],
            ["Orders and finance", "Order, OrderItem, OrderStatusHistory, Invoice, InvoiceLineItem, CreditNote, InvoiceSequence"],
            ["Returns and fulfillment", "Return, Refund, Shipment"],
            ["Trust and configuration", "Review, PaymentEvent, AdminActivityLog, RolePermission, StoreSetting, StoreSettingsAudit"],
        ],
        [1.7, 4.8],
    )

    doc.add_heading("API footprint", level=2)
    add_paragraph(
        doc,
        "The checked-in endpoint layer spans 15 domain files and 169 explicit handler decorators. Router aggregation in backend/app/api/v1/router.py groups the surface into authentication, admin 2FA, catalog, cart and coupons, orders, payments, invoices, shipping and returns, dashboard and reports, privacy, user profile, images, products, and store settings.",
        after=8,
    )
    add_table(
        doc,
        ["Key endpoint modules", "Primary responsibilities"],
        [
            ["auth.py, totp.py, user.py, privacy.py", "Session lifecycle, admin 2FA, profile and address CRUD, consent and deletion workflows"],
            ["catalog.py, wishlist_reviews.py, cart_coupons.py", "Public browsing, product search and filters, wishlist, reviews, cart mutation, coupon validation"],
            ["orders.py, payments.py, invoices.py", "Checkout summary, order placement, order admin actions, payment creation or verification, invoice and credit note access"],
            ["shipping_returns.py", "Shipment updates, returns approval flow, refund initiation"],
            ["admin_products.py, admin_images.py, admin_dashboard.py, admin_settings.py", "Product and inventory admin, image pipeline, dashboards and reports, store configuration"],
        ],
        [2.2, 4.3],
    )

    doc.add_heading("3. Frontend Architecture", level=1)
    add_paragraph(
        doc,
        "The frontend is a React 19 single-page application built with Vite 7 and Tailwind CSS v4. App.jsx restores sessions via httpOnly cookies on boot, mounts customer and admin route trees, and uses ProtectedRoute plus AdminRoute wrappers to gate access.",
        after=8,
    )
    add_bullets(
        doc,
        [
            "The customer experience includes landing, category, product listing and detail, search, cart, checkout, profile, privacy, orders, and invoice viewing pages.",
            "The admin experience includes dashboard, product management, product form, category and attribute managers, coupons, inventory, orders, returns, invoices, reports, users, audit logs, and settings.",
            "State management is split across nine Zustand stores, including auth, cart, catalog, order, UI, toast, address, and admin-specific state containers.",
            "The API layer is currently concentrated in apiClient.js and adminApi.js rather than the larger multi-file API split described in the older reference PDF.",
        ],
    )
    add_table(
        doc,
        ["Frontend layer", "Current implementation details"],
        [
            ["Routing", "React Router tree in App.jsx with lazy loading for lower-frequency pages and a redirect from /admin/login to /login"],
            ["Auth transport", "Cookie-based auth with withCredentials enabled, 401 refresh retry logic, and auth expiry event broadcast from the Axios interceptor"],
            ["Component organization", "Shared UI, auth guards, admin layouts, checkout widgets, catalog components, and reusable common controls"],
            ["Testing", "Playwright config auto-starts npm run dev on port 3000; vitest is installed but package.json still lacks a canonical test script"],
        ],
        [1.8, 4.7],
    )

    doc.add_heading("4. Delivery, Testing, and Operations", level=1)
    add_bullets(
        doc,
        [
            "Local infra: docker compose up postgres redis api from repo root, with backend/.env loaded by the api service.",
            "Frontend commands: npm run dev, npm run build, npm run lint, and npm run preview from frontend/.",
            "Backend test configuration: pytest settings live in backend/pytest.ini.toml, although testpaths is still set to tests while the checked-in suite lives under backend/app/tests.",
            "The deployment workflow builds backend/Dockerfile.lambda into ECR, updates the Lambda function image, builds the frontend on Node 20, syncs dist/ to S3, and invalidates CloudFront.",
        ],
    )

    doc.add_heading("5. Current Drift From The March 2026 Reference PDF", level=1)
    add_bullets(
        doc,
        [
            "The reference PDF cites 172 routes; the current endpoint files expose 169 explicit handler decorators under backend/app/api/v1/endpoints.",
            "The reference PDF lists 27 or 30 schema tables in places; the live models.py currently declares 33 ORM table classes.",
            "The old document describes many dedicated frontend API modules such as authApi.js and catalogApi.js, but the current repo mainly exposes apiClient.js plus a consolidated adminApi.js.",
            "The March document says app startup wires reservation expiry and FX sync together; current app/main.py starts only reservation expiry, while FX sync and deletion processing are reachable through lambda_handler scheduled task routing.",
            "Repository instructions and current workflow differ on deploy branch expectations; the checked-in GitHub Actions file deploy-dev-v3.yml presently triggers on pushes to develop.",
            "The codebase still contains some transitional artifacts such as frontend/src/pages/customer/index.jsx and legacy route helper files that are not the main runtime entrypoints.",
        ],
    )

    doc.add_heading("6. Summary", level=1)
    add_paragraph(
        doc,
        "Ashmiwebportal remains a fairly broad commerce platform with customer shopping flows, admin operations, GST-aware invoicing, dual payment gateway support, and a mixed local-plus-AWS deployment story. "
        "The architecture is still recognizable from the earlier summary, but the live repository now differs enough in counts, file layout, and workflow wiring that the March 2026 PDF should be treated as historical context rather than the authoritative description.",
        after=8,
    )
    add_paragraph(
        doc,
        "Method: this summary was generated from the current checked-in source tree, including backend and frontend entrypoints, package manifests, workflow YAML, test configuration, and direct file counts measured on 14 May 2026.",
        size=10,
        color=MUTED,
        italic=True,
        after=0,
    )

    doc.save(out_path)


def main() -> None:
    output_dir = ROOT / "generated"
    output_dir.mkdir(exist_ok=True)
    docx_path = output_dir / "Ashmiwebportal_Codebase_Summary_2026-05-14.docx"
    metrics = collect_metrics()
    build_doc(metrics, docx_path)
    print(docx_path)


if __name__ == "__main__":
    main()
